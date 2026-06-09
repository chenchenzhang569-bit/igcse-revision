#!/usr/bin/env python3
"""生成 IGMaster 完整配置报告 Word 文档"""

from docx import Document
from docx.shared import Inches, Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

doc = Document()

# ─── 样式 ───
style = doc.styles['Normal']
font = style.font
font.name = '微软雅黑'
font.size = Pt(10.5)
style.element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def add_table(doc, headers, rows):
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        cell.text = h
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(9)
    # Rows
    for row_data in rows:
        row = table.add_row()
        for i, val in enumerate(row_data):
            row.cells[i].text = str(val)
            for p in row.cells[i].paragraphs:
                for r in p.runs:
                    r.font.size = Pt(9)
    doc.add_paragraph()
    return table

# ══════════════════════════════════
# 封面
# ══════════════════════════════════
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('IGMaster 完整配置报告')
run.font.size = Pt(28)
run.bold = True

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run(f'IGCSE 复习平台 · 安全 · 备份 · 部署 · 配置')
run.font.size = Pt(14)
run.font.color.rgb = RGBColor(0x00, 0x1C, 0x71)

doc.add_paragraph()
date_para = doc.add_paragraph()
date_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
date_para.add_run(f'生成日期: {datetime.date.today().strftime("%Y年%m月%d日")}').font.size = Pt(11)

doc.add_paragraph()
ver = doc.add_paragraph()
ver.alignment = WD_ALIGN_PARAGRAPH.CENTER
ver.add_run('版本: 1.0').font.size = Pt(10)

doc.add_page_break()

# ══════════════════════════════════
# 目录
# ══════════════════════════════════
doc.add_heading('目录', level=1)
toc_items = [
    '一、项目概览',
    '二、分支与部署工作流',
    '三、安全审计与修复',
    '四、数据库 RLS 策略',
    '五、备份方案',
    '六、环境变量清单',
    '七、回滚方案',
    '八、Storage 存储桶',
]
for item in toc_items:
    p = doc.add_paragraph(item)
    p.paragraph_format.space_after = Pt(4)

doc.add_page_break()

# ══════════════════════════════════
# 一、项目概览
# ══════════════════════════════════
doc.add_heading('一、项目概览', level=1)

doc.add_heading('基本信息', level=2)
add_table(doc, ['项目', '值'], [
    ['项目名称', 'IGMaster (igcse-revision)'],
    ['网站域名', 'igmaster.org'],
    ['技术栈', 'Next.js 14+ / Supabase / Tailwind CSS'],
    ['托管平台', 'Vercel (dev) + Zeabur 香港节点 (production)'],
    ['数据库', 'Supabase PostgreSQL (Pro 套餐)'],
    ['文件存储', 'Supabase Storage'],
    ['邮箱服务', 'Resend (免费 3,000 封/月)'],
    ['支付', '支付宝（搁置中，待注册个体户）'],
    ['GitHub', 'github.com/chenchenzhang569-bit/igcse-revision'],
])

doc.add_heading('服务器位置', level=2)
add_table(doc, ['平台', '节点/区域', '用途', '费用'],
    [
        ['Vercel', '全球 CDN', '开发/测试 (main 分支)', 'Hobby (免费)'],
        ['Zeabur', '香港节点', '正式上线 (production 分支)', 'Developer ($5/月)'],
        ['Supabase', '美国 (默认)', '数据库 + Storage + Auth', 'Pro ($25/月)'],
        ['阿里云', '中国大陆', '域名 DNS 解析', '按量付费'],
    ])

doc.add_page_break()

# ══════════════════════════════════
# 二、分支与部署工作流
# ══════════════════════════════════
doc.add_heading('二、分支与部署工作流', level=1)

doc.add_heading('分支模型', level=2)
add_table(doc, ['分支', '部署到', '用途', '保护规则'],
    [
        ['main', 'Vercel', '开发/测试', '无'],
        ['production', 'Zeabur 香港', '正式上线', 'PR + 1 approve + 管理员'],
    ])

doc.add_heading('工作流程', level=2)
p = doc.add_paragraph()
p.add_run('开发 → 推 main → Vercel 自动部署\n'
           '            → 测试通过\n'
           '            → 开 PR (main → production)\n'
           '            → 用户 GitHub approve\n'
           '            → 合并 → Zeabur 自动部署').font.size = Pt(10)

doc.add_heading('注意事项', level=2)
notes = [
    '推 main 前必须先问用户（微信说"推"才能推）',
    '从不合自己合并 production，必须走 PR + approve',
    'GitHub 默认分支设为 production（Zeabur 自动跟随）',
    'Vercel 跟随 GitHub 默认分支，自动把 production 当生产分支',
]
for n in notes:
    doc.add_paragraph(n, style='List Bullet')

doc.add_page_break()

# ══════════════════════════════════
# 三、安全审计与修复
# ══════════════════════════════════
doc.add_heading('三、安全审计与修复', level=1)
p = doc.add_paragraph()
p.add_run('审计日期: 2026-06-02').font.color.rgb = RGBColor(100, 100, 100)

doc.add_heading('已修复问题', level=2)
add_table(doc, ['文件', '问题', '修复'],
    [
        ['errors/route.ts', '无 admin 验证，任意登录用户可查看/修改用户报错', '添加 requireAdmin() 验证'],
        ['download-sme-ms/route.ts', '无 admin 验证 + SME 密码硬编码在源码', '添加 requireAdmin() + 改为环境变量'],
        ['.env.example', '缺少 SME 环境变量说明', '添加 SME_EMAIL / SME_PASSWORD 条目'],
    ])

doc.add_heading('已验证的安全项', level=2)
items = [
    '✅ 代码中无硬编码密钥（API key、数据库密码等）',
    '✅ Service role key 仅服务端使用（admin API routes）',
    '✅ 付费内容通过 access.ts + cookie auth 服务端验证',
    '✅ 中间件保护非公开页面，未登录重定向到 /login',
    '✅ 所有 admin 路由都已验证管理员身份',
]
for item in items:
    doc.add_paragraph(item)

doc.add_page_break()

# ══════════════════════════════════
# 四、数据库 RLS 策略
# ══════════════════════════════════
doc.add_heading('四、数据库 RLS 策略', level=1)
p = doc.add_paragraph()
p.add_run('所有表均已开启 Row Level Security (RLS)。以下是通过 Supabase REST API 验证的结果：')

doc.add_heading('表级 RLS', level=2)
add_table(doc, ['表名', '策略', '说明'],
    [
        ['exam_boards', '公开读', '所有人都能查询考试局列表'],
        ['subjects', '公开读', '所有人都能查询科目'],
        ['topics', '公开读', '所有人都能查询主题'],
        ['subtopics', '公开读', '所有人都能查询子主题'],
        ['questions', '公开读', '所有人都能查询题目'],
        ['notes', '公开读', '所有人都能查询笔记'],
        ['past_papers', '公开读', '所有人都能查询试卷'],
        ['purchases', '仅自己读', '只能看到自己的购买记录 ✅'],
        ['profiles', '仅自己读', '只能看到自己的个人信息 ✅'],
        ['error_reports', '管理员读', '（有递归问题，但 admin 用 service_role 绕过）'],
        ['login_events', '管理员读', '仅 admin 可访问'],
        ['user_roles', '管理员读', '仅 admin 可访问'],
        ['app_config', '管理员读', '存储 Resend API key 等配置'],
        ['mock_exam_*', '公开读', '模拟考数据公开'],
    ])

doc.add_heading('Storage 桶 RLS', level=2)
add_table(doc, ['桶名', '公开', '说明'],
    [
        ['past-papers', '是', '学生需要下载 PDF 做题'],
        ['mock-exams', '是', '模拟考试卷 PDF'],
        ['notes-pdfs', '是', '笔记 PDF'],
        ['sme-images', '是', 'SME 截图图片'],
        ['sme-raw-backup', '否', '原始数据备份（私有）'],
        ['scripts', '是', '脚本文件'],
    ])

doc.add_page_break()

# ══════════════════════════════════
# 五、备份方案
# ══════════════════════════════════
doc.add_heading('五、备份方案', level=1)

doc.add_heading('定时备份任务', level=2)
add_table(doc, ['备份类型', '频率', '时间', '保存周期', '位置'],
    [
        ['数据库导出', '每天', '凌晨 3:00', '30 天', '~/backups/db/'],
        ['Storage 文件', '每周', '周日 4:00', '2 份', '~/backups/storage/'],
        ['代码仓库', '每周', '周日 4:00', '7 天', '~/backups/code/'],
    ])

doc.add_heading('备份脚本', level=2)
p = doc.add_paragraph()
p.add_run('主脚本: ').bold = True
p.add_run('scripts/backup_all.py')
doc.add_paragraph()

add_table(doc, ['命令', '说明'],
    [
        ['python3 scripts/backup_all.py', '完整备份（DB + Storage + 代码）'],
        ['python3 scripts/backup_all.py --type db', '仅数据库'],
        ['python3 scripts/backup_all.py --type storage', '仅 Storage 文件'],
        ['python3 scripts/backup_all.py --type code', '仅代码仓库'],
    ])

doc.add_heading('备份文件结构', level=2)
p = doc.add_paragraph()
p.add_run('~/backups/').bold = True
doc.add_paragraph('  ├── db/                          # 数据库备份（每日）')
doc.add_paragraph('  │   ├── db_20260603_003332.json.gz    ← 最新')
doc.add_paragraph('  │   └── ...')
doc.add_paragraph('  ├── storage/                     # Storage 文件备份（每周）')
doc.add_paragraph('  │   └── 20260603/')
doc.add_paragraph('  │       ├── past-papers/')
doc.add_paragraph('  │       ├── mock-exams/')
doc.add_paragraph('  │       ├── notes-pdfs/')
doc.add_paragraph('  │       ├── sme-raw-backup/')
doc.add_paragraph('  │       ├── sme-images/')
doc.add_paragraph('  │       └── scripts/')
doc.add_paragraph('  └── code/                        # 代码备份（每周）')
doc.add_paragraph('      └── code_20260603_*.bundle')

doc.add_heading('备份内容详情', level=2)
p = doc.add_paragraph()
p.add_run('数据库备份包含以下 16 张表 + Auth 用户（首次测试 12,422 行，约 40.5MB）：').font.size = Pt(10)
tables_list = [
    ('exam_boards', '2 行'), ('subjects', '11 行'), ('topics', '143 行'),
    ('subtopics', '390 行'), ('notes', '450 行'), ('questions', '5,379 行'),
    ('past_papers', '4,137 行'), ('purchases', '21 行'), ('profiles', '17 行'),
    ('error_reports', '10 行'), ('login_events', '13 行'), ('user_roles', '1 行'),
    ('app_config', '3 行'), ('mock_exam_sets', '49 行'),
    ('mock_exam_papers', '97 行'), ('mock_exam_questions', '1,699 行'),
    ('auth_users', '17 人'),
]
add_table(doc, ['表名', '数据量'], tables_list)

doc.add_page_break()

# ══════════════════════════════════
# 六、环境变量清单
# ══════════════════════════════════
doc.add_heading('六、环境变量清单', level=1)
p = doc.add_paragraph()
p.add_run('需要在 Vercel 和 Zeabur 两个平台都设置的变量。').font.size = Pt(10)

doc.add_heading('必须设置的变量', level=2)
add_table(doc, ['变量名', '说明', '示例'],
    [
        ['NEXT_PUBLIC_SUPABASE_URL', 'Supabase 项目 URL', 'https://aondldqwwvttwpervrfq.supabase.co'],
        ['NEXT_PUBLIC_SUPABASE_ANON_KEY', 'Supabase 公开 anon key', "os.environ.get("SERVICE_ROLE_KEY", "")"],
        ['SUPABASE_SERVICE_ROLE_KEY', '管理密钥（最高权限，勿泄露）', "os.environ.get("SERVICE_ROLE_KEY", "")"],
        ['NEXT_PUBLIC_SITE_URL', '网站地址', 'https://igmaster.org'],
        ['SME_EMAIL', 'SaveMyExams 登录邮箱', 'inspiringchermann@vmail.dev'],
        ['SME_PASSWORD', 'SaveMyExams 登录密码', '********'],
    ])

doc.add_heading('支付宝相关（暂时搁置）', level=2)
add_table(doc, ['变量名', '说明'],
    [
        ['ALIPAY_APP_ID', '支付宝 APPID'],
        ['ALIPAY_PRIVATE_KEY', '支付宝 RSA 私钥'],
        ['ALIPAY_PUBLIC_KEY', '支付宝 RSA 公钥'],
        ['ALIPAY_GATEWAY', '支付宝网关 URL（默认沙箱）'],
        ['ALIPAY_NOTIFY_URL', '异步通知 URL'],
        ['ALIPAY_RETURN_URL', '同步跳转 URL'],
    ])

doc.add_page_break()

# ══════════════════════════════════
# 七、回滚方案
# ══════════════════════════════════
doc.add_heading('七、回滚方案', level=1)
p = doc.add_paragraph()
p.add_run('详细回滚步骤见 docs/rollback-plan.md。以下是快速参考：')

doc.add_heading('代码回滚', level=2)
steps = [
    '查看最近提交: git log --oneline -10',
    '回滚到上一版本: git revert HEAD',
    '推送到 main: git push origin main',
    '推送 production 同理',
]
for s in steps:
    doc.add_paragraph(s, style='List Bullet')

doc.add_heading('数据库恢复', level=2)
p = doc.add_paragraph()
p.add_run('找到备份文件解压后，通过 Supabase REST API 逐表恢复。')
doc.add_paragraph('备份位置: ~/backups/db/db_最新日期.json.gz', style='List Bullet')
doc.add_paragraph('恢复脚本: 参考 docs/rollback-plan.md 中的恢复命令', style='List Bullet')

doc.add_heading('域名切换（Zeabur 宕机时）', level=2)
p = doc.add_paragraph()
p.add_run('登录阿里云 DNS (dns.console.aliyun.com)，将 CNAME 从 ')
p.add_run('igmaster.zeabur.app').bold = True
p.add_run(' 指向 Vercel 域名。')

doc.add_page_break()

# ══════════════════════════════════
# 八、Storage 存储桶
# ══════════════════════════════════
doc.add_heading('八、Storage 存储桶', level=1)

add_table(doc, ['桶名', 'Public', '存储内容'],
    [
        ['past-papers', '是', '历年真题 PDF'],
        ['mock-exams', '是', '模拟考试卷 PDF'],
        ['notes-pdfs', '是', '笔记 PDF 文件'],
        ['sme-raw-backup', '否', 'SME 原始数据 JSON 备份'],
        ['sme-images', '是', 'SME 截图/图片'],
        ['scripts', '是', '辅助脚本和配置文件'],
    ])

p = doc.add_paragraph()
p.add_run('\n所有 Public 桶的文件均可通过以下格式直接访问（无需认证）：').font.size = Pt(10)
doc.add_paragraph('https://aondldqwwvttwpervrfq.supabase.co/storage/v1/object/public/{桶名}/{文件路径}')

doc.add_page_break()

# ══════════════════════════════════
# 附录
# ══════════════════════════════════
doc.add_heading('附录：技术债 / 待办事项', level=1)

add_table(doc, ['优先级', '事项', '说明'],
    [
        ['高', '统一 admin 验证体系', 'requireAdmin() vs checkAdmin() 两套并存'],
        ['中', '修复 error_reports RLS 递归问题', '当前 anon key 读取时报错，但不影响 admin 功能'],
        ['中', '合并 main → production', 'auth 安全修复已推 main 但 Zeabur 未部署'],
        ['低', '支付宝支付上线', '需注册个体户后配置'],
        ['低', 'SMTP SPF 记录配置', '需阿里云 AccessKey'],
        ['低', 'Vercel 绑定 igmaster.org 域名', '域名切换 Vercel 接管时需要'],
    ])

# 保存
doc.save('/home/ubuntu/igcse-revision/IGMaster_配置报告.docx')
print("✅ Word 文档生成成功: IGMaster_配置报告.docx")
